"""Generate the requirements analysis Word report."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

def add_para(text, bold=False, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=len(rows)+1, cols=len(headers), style='Table Grid')
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()
    return table

# ═══════════════ Title Page ═══════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('实验室智能签到管理系统\n需求分析报告')
run.bold = True
run.font.size = Pt(24)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run(f'文档版本：V1.0\n').font.size = Pt(12)
info.add_run(f'编制日期：{datetime.now().strftime("%Y年%m月%d日")}\n').font.size = Pt(12)
info.add_run('编制人：郭晓飞 项目开发团队\n').font.size = Pt(12)
info.add_run('审核人：________（指导教师）\n').font.size = Pt(12)
info.add_run('批准人：________（实验室负责人）').font.size = Pt(12)

doc.add_page_break()

# ═══════════════ Revision History ═══════════════
doc.add_heading('修订记录', level=2)
add_table(
    ['版本', '修订日期', '修订内容', '修订人'],
    [['V1.0', datetime.now().strftime('%Y-%m-%d'), '初始版本', '郭晓飞']]
)

# ═══════════════ 1. Introduction ═══════════════
doc.add_heading('1. 引言', level=2)

doc.add_heading('1.1 编写目的', level=3)
add_para('本文档为"实验室智能签到管理系统"的正式需求分析说明书。文档全面、准确地描述了本项目的目标、功能需求、非功能需求及运行环境要求，为项目的设计、开发、测试和验收提供依据。', indent=True)

doc.add_heading('1.2 项目背景', level=3)
add_para('高校实验室是开展实践教学和科学研究的重要场所。传统的实验室考勤管理存在以下突出问题：', indent=True)
add_para('(1) 签到效率低下——纸质签到或口头点名耗费大量课堂时间。')
add_para('(2) 代签现象严重——传统方式难以杜绝代替他人签到。')
add_para('(3) 安全管理薄弱——无法实时掌握实验室内人员情况。')
add_para('(4) 统计工作繁琐——学期末人工汇总签到数据工作量大。')
add_para('(5) 缺乏智能化——实验室使用情况无法自动采集和分析。')
add_para('为解决上述痛点，本项目拟开发一套融合二维码扫码、GPS定位、人脸识别等技术的实验室智能签到管理系统。', indent=True)

doc.add_heading('1.3 项目范围', level=3)
doc.add_heading('1.3.1 功能范围', level=4)
add_para('本次开发实现以下核心功能模块：')
add_para('1) 手机扫码签到模块：基于QR二维码+高德GPS定位（<=100米）+姓名确认的多因素签到')
add_para('2) 摄像头人脸识别签到模块：利用摄像头自动识别进出人员并完成签到/签退')
add_para('3) 自动签退模块：支持超时自动签退、超距自动签退、新地点签到自动签退三种触发机制')
add_para('4) 角色与权限管理模块：设置管理员、指导教师、学生三级角色')
add_para('5) 统计与管理模块：签到记录查询、出勤率统计、人员管理、签到点管理')
add_para('6) 视频批量签到模块：上传视频自动识别所有人员并批量完成签到')

doc.add_heading('1.3.2 不包括的功能范围', level=4)
add_para('以下功能暂不纳入本次开发范围：')
add_para('  - 与学校教务系统的数据对接')
add_para('  - 高精度室内定位（UWB/蓝牙信标硬件方案）')
add_para('  - 移动端原生App（首版采用Web App）')
add_para('  - 与校园一卡通/门禁系统的硬件集成')

doc.add_heading('1.4 术语与定义', level=3)
add_table(
    ['术语/缩写', '全称与定义'],
    [
        ['GCJ-02', '国家测绘局制定的地理坐标加密标准（火星坐标系），高德地图使用此坐标系。'],
        ['WGS-84', 'World Geodetic System 1984，GPS设备返回的原始坐标系统，国内与GCJ-02偏差100-500米。'],
        ['QR码', 'Quick Response Code，本系统中用于承载签到/签退URL。'],
        ['cpolar', '内网穿透工具，将本地服务映射到公网HTTPS域名。'],
        ['embedding', '人脸特征向量（512维），用于人脸比对。'],
        ['Haversine', '球面距离计算公式，用于计算用户位置与签到点的地表距离。'],
        ['SCRFD', '轻量级人脸检测模型。'],
        ['ONNX', '开放神经网络模型交换格式。'],
    ]
)

doc.add_heading('1.5 参考文档', level=3)
add_para('  - 《2024-2025学年实验室管理规范》')
add_para('  - 《高德开放平台Web JS API开发文档》')
add_para('  - 《InsightFace模型文档》（github.com/deepinsight/insightface）')
add_para('  - 前期需求沟通会议纪要（2026年6月）')

# ═══════════════ 2. Overview ═══════════════
doc.add_heading('2. 产品需求概述', level=2)

doc.add_heading('2.1 产品愿景', level=3)
add_para('打造一套低成本、易部署、高精度的实验室智能签到系统，融合QR扫码、GPS定位与人脸识别技术，实现"扫码即签到、摄像头无感签到、离场自动签退"的全流程自动化管理。', indent=True)

doc.add_heading('2.2 项目目标', level=3)
add_para('业务目标：', bold=True)
add_para('  - 将单次签到时间从2-3分钟缩短至10秒以内')
add_para('  - 将代签率降至0%')
add_para('  - 学期末考勤统计时间缩短至1分钟')
add_para('用户目标：', bold=True)
add_para('  - 学生无需下载App，手机扫码或摄像头即可签到')
add_para('  - 教师无需手动点名，系统自动完成考勤')
add_para('  - 管理员1分钟内完成新增用户和签到点配置')
add_para('技术目标：', bold=True)
add_para('  - GPS定位精度目标<=50米（室内），最大允许偏差100米')
add_para('  - 人脸识别准确率>=95%')
add_para('  - 系统支持50人以上同时签到，响应时间<=3秒')

doc.add_heading('2.3 用户及特征', level=3)
add_table(
    ['用户角色', '职责描述', '使用频率', '技术水平', '核心需求与痛点'],
    [
        ['实验室管理员', '管理系统用户、签到点、查看统计报表', '每日多次', '中等', '快速查看实时出勤；便捷的人员管理'],
        ['实验指导教师', '生成签到二维码、查看签到情况', '每次实验课', '中等', '课前快速生成签到码；课上实时查看进度'],
        ['学生', '通过手机扫码完成签到/签退', '每次实验课', '基本手机操作', '签到流程简单快速；不下载App'],
    ]
)

# ═══════════════ 3. Detailed Requirements ═══════════════
doc.add_heading('3. 详细功能需求', level=2)

doc.add_heading('3.1 功能模块概览', level=3)
add_para('1. 用户认证与角色权限模块     2. QR码签到模块')
add_para('3. 摄像头人脸签到模块         4. 自动签退模块')
add_para('5. 签到点管理模块             6. 统计与报表模块')
add_para('7. 视频批量处理模块')

doc.add_heading('3.2 模块一：用户认证与角色权限', level=3)
add_para('（1）用户登录：用户名密码登录，bcrypt验证，生成JWT令牌（24h有效期）。登录后按角色自动跳转对应页面。')
add_para('（2）角色权限：管理员->全部功能+全部数据；教师->QR生成+人员管理+统计；学生->个人签到状态+历史记录。')
add_para('（3）用户管理：管理员可创建/禁用用户，为学生注册人脸照片。')

doc.add_heading('3.3 模块二：QR码扫码签到', level=3)
add_para('（1）生成二维码：教师选择签到类型和签到点，系统生成含UUID令牌的QR码（5分钟有效，一次性使用）。')
add_para('（2）GPS定位校验：手机GPS获取WGS-84坐标->后端转换为GCJ-02->Haversine公式计算距离-><=100米通过。')
add_para('（3）姓名确认签到：定位通过后输入姓名->匹配用户->检查无活跃签到->创建签到记录。')

doc.add_heading('3.4 模块三：摄像头人脸识别签到', level=3)
add_para('（1）人脸检测与识别：InsightFace SCRFD-10GF检测+ResNet50提取512维特征向量+余弦相似度匹配（阈值0.25）。')
add_para('（2）视频批量签到：上传视频->每0.5秒采样一帧->检测+去重（阈值0.45）->匹配用户库->批量签到。')

doc.add_heading('3.5 模块四：自动签退机制', level=3)
add_para('（1）超时自动签退：后台每5分钟扫描，签到超过4小时自动签退。')
add_para('（2）超距自动签退（预留）：离开签到点超过设定范围触发。')
add_para('（3）新签到自动签退：在新点签到时自动完成之前的签到记录。')
add_para('另支持手动签退：扫描签退QR码+输入姓名。')

doc.add_heading('3.6 模块五~七', level=3)
add_para('签到点管理：创建/停用签到点，设置名称、GCJ-02坐标、有效范围半径。')
add_para('统计与报表：实时仪表盘（总人数/已签到/未签到/平均时长）、签到记录查询、出勤率统计。')
add_para('前端界面：登录页、管理后台（6个标签页）、学生面板（签到状态+历史）、扫码签到页（3步骤）。')

# ═══════════════ 4. Non-functional ═══════════════
doc.add_heading('4. 非功能需求', level=2)

doc.add_heading('4.1 性能需求', level=3)
add_para('  - 普通页面响应时间<=2秒')
add_para('  - 单张人脸识别<=300ms（GPU）')
add_para('  - 支持50人同时在线')
add_para('  - 支持存储5年签到历史（约10万条）')

doc.add_heading('4.2 安全需求', level=3)
add_para('  - 密码bcrypt哈希存储，禁止明文')
add_para('  - JWT令牌24h有效期')
add_para('  - 严格角色权限控制')
add_para('  - API密钥存.env环境变量')
add_para('  - 防护SQL注入、XSS、CSRF')

doc.add_heading('4.3 可靠性需求', level=3)
add_para('  - 7x24小时运行')
add_para('  - 数据库崩溃后2小时内恢复')
add_para('  - 人脸模型故障时姓名签到不受影响')

doc.add_heading('4.4 易用性需求', level=3)
add_para('  - 学生签到<=3步操作')
add_para('  - 管理员1分钟内完成基础配置')
add_para('  - 中文界面+清晰错误提示')
add_para('  - 移动端适配主流手机屏幕')

doc.add_heading('4.5 可维护性需求', level=3)
add_para('  - 模块化设计，路由/服务/数据层分离')
add_para('  - 遵循PEP 8规范')
add_para('  - 配置集中管理（config.py+环境变量）')
add_para('  - ORM框架支持数据库切换')

# ═══════════════ 5. Environment ═══════════════
doc.add_heading('5. 运行环境需求', level=2)

doc.add_heading('5.1 服务器环境', level=3)
add_table(
    ['项目', '最低配置', '推荐配置'],
    [
        ['CPU', '4核', '8核以上'],
        ['内存', '8GB', '16GB以上'],
        ['显卡', 'GTX 1060 (6GB)', 'RTX 3060 (8GB)'],
        ['操作系统', 'Win10/11 或 Ubuntu 20.04+', 'Win11 或 Ubuntu 22.04+'],
        ['Python', '3.10', '3.10'],
        ['CUDA', '12.1+', '12.7+'],
    ]
)

doc.add_heading('5.2 客户端环境', level=3)
add_para('PC端：Chrome 90+/Edge 90+/Firefox 88+，分辨率>=1366x768。')
add_para('移动端：iOS 14+ Safari / Android 10+ Chrome，需GPS模块。')

doc.add_heading('5.3 网络环境', level=3)
add_para('  - 服务器通过cpolar隧道获得公网HTTPS域名')
add_para('  - 局域网带宽>=100Mbps')
add_para('  - 服务端口：8080（可配置）')

doc.add_heading('5.4 第三方服务依赖', level=3)
add_table(
    ['服务', '用途', '免费额度'],
    [
        ['高德Web JS API', '手机GPS定位', '每日30万次'],
        ['高德REST API', '逆地理编码', '每日30万次'],
        ['cpolar', 'HTTPS隧道', '免费版1隧道'],
        ['InsightFace (buffalo_l)', '人脸检测+识别(离线)', '完全免费'],
    ]
)

# ═══════════════ 6. Verification ═══════════════
doc.add_heading('6. 需求确认与验证', level=2)

add_para('6.1 确认方式', bold=True)
add_para('通过需求评审会议、原型系统演示、功能点逐条确认。')

add_para('6.2 验收标准', bold=True)
add_para('(1) 功能验收：7个功能模块核心功能点在真实环境可正常运行。')
add_para('(2) 定位精度验收：至少3台不同型号手机测试，偏差均<=100米。')
add_para('(3) 人脸识别验收：正常光照下已注册人员识别准确率>=95%。')
add_para('(4) 并发验收：10台设备同时签到，30秒内完成，系统无错误。')
add_para('(5) 容错验收：故障场景下系统给出明确提示。')

add_para('6.3 签字确认', bold=True)
add_table(
    ['角色', '姓名', '签字', '日期'],
    [
        ['项目负责人', '', '', ''],
        ['指导教师', '', '', ''],
        ['实验室管理员代表', '', '', ''],
        ['学生代表', '', '', ''],
    ]
)

# ═══════════════ 7. Appendices ═══════════════
doc.add_heading('7. 附录', level=2)

doc.add_heading('附录1：系统架构', level=3)
add_para('前端层：原生HTML/CSS/JS（响应式设计）')
add_para('后端层：Python FastAPI（异步RESTful API）')
add_para('数据层：SQLite + SQLAlchemy ORM')
add_para('AI推理层：InsightFace + ONNX Runtime GPU')
add_para('外部服务：高德API、cpolar隧道')

doc.add_heading('附录2：核心业务流程', level=3)
add_para('签到流程：管理员生成QR码->学生扫码->定位校验->输入姓名->匹配用户->创建记录->返回成功。')
add_para('签退流程：扫码/超时/新签到触发->查找活跃记录->写入签退时间->更新状态。')
add_para('视频签到流程：上传视频->逐帧检测->跨帧去重->匹配人脸库->批量创建记录。')

doc.add_heading('附录3：数据库核心表', level=3)
add_table(
    ['表名', '核心字段', '说明'],
    [
        ['users', 'id, username, name, role, face_embedding, password_hash', '用户表'],
        ['locations', 'id, name, latitude, longitude, radius_meters', '签到点表'],
        ['checkins', 'id, user_id, check_in_time, check_out_time, status', '签到记录表'],
        ['qr_sessions', 'id, token, type, location_id, expires_at', 'QR会话表'],
    ]
)

doc.add_heading('附录4：技术选型', level=3)
add_table(
    ['层次', '技术栈', '理由'],
    [
        ['后端框架', 'FastAPI (Python 3.10 async)', '异步高性能，自动API文档'],
        ['数据库', 'SQLite + SQLAlchemy 2.0', '零配置，ORM可切换'],
        ['人脸检测', 'InsightFace SCRFD-10GF', '高精度，GPU加速'],
        ['人脸识别', 'InsightFace ResNet50@WebFace600K', '512维embedding'],
        ['推理引擎', 'ONNX Runtime GPU', '跨平台GPU推理'],
        ['坐标转换', 'WGS-84->GCJ-02数学模型', '毫秒级转换'],
        ['距离计算', 'Haversine公式', '100米级精确'],
        ['前端', '原生HTML/CSS/JS', '无框架，移动端兼容好'],
        ['认证', 'JWT + bcrypt', '无状态，安全哈希'],
        ['HTTPS', 'cpolar', '一键公网HTTPS'],
    ]
)

out_path = r'D:\insightface\checkin_system\实验室智能签到管理系统_需求分析报告.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
